import pytest
from django.core.files.base import ContentFile
from io import BytesIO
from pypdf import PdfWriter
from reportlab.pdfgen import canvas
from docx import Document as DocxDocument

from knowledge.extractors import (
    DOCXExtractor,
    DocumentExtractorFactory,
    PDFExtractor,
    PlainTextExtractor,
    UnsupportedDocumentTypeError,
)
from knowledge.models import Document


@pytest.mark.django_db
def test_plain_text_extractor_extracts_document_text(
    django_user_model,
):
    user = django_user_model.objects.create_user(
        username="extractuser",
        password="test-password-123",
    )

    document = Document.objects.create(
        user=user,
        title="Text Document",
        original_filename="notes.txt",
        content_type="text/plain",
    )

    document.file.save(
        "notes.txt",
        ContentFile(
            b"Machine learning is a field of AI."
        ),
    )

    extractor = PlainTextExtractor()

    text = extractor.extract(document)

    assert text == "Machine learning is a field of AI."


@pytest.mark.django_db
def test_plain_text_extractor_rejects_invalid_utf8(
    django_user_model,
):
    user = django_user_model.objects.create_user(
        username="invalidtextuser",
        password="test-password-123",
    )

    document = Document.objects.create(
        user=user,
        title="Invalid Text",
        original_filename="invalid.txt",
        content_type="text/plain",
    )

    document.file.save(
        "invalid.txt",
        ContentFile(
            b"Valid text\xff\xfe"
        ),
    )

    extractor = PlainTextExtractor()

    with pytest.raises(UnicodeDecodeError):
        extractor.extract(document)

def test_extractor_factory_returns_plain_text_extractor():
    extractor = DocumentExtractorFactory.get_extractor(
        "text/plain"
    )

    assert isinstance(
        extractor,
        PlainTextExtractor,
    )

def test_extractor_factory_rejects_unsupported_type():
    with pytest.raises(UnsupportedDocumentTypeError):
        DocumentExtractorFactory.get_extractor(
            "application/octet-stream"
        )

@pytest.mark.django_db
def test_pdf_extractor_extracts_text(
    django_user_model,
):
    user = django_user_model.objects.create_user(
        username="pdfextractuser",
        password="test-password-123",
    )

    pdf_buffer = BytesIO()

    pdf = canvas.Canvas(pdf_buffer)
    pdf.drawString(
        100,
        750,
        "Machine learning is useful.",
    )
    pdf.save()

    document = Document.objects.create(
        user=user,
        title="PDF Document",
        original_filename="document.pdf",
        content_type="application/pdf",
    )

    document.file.save(
        "document.pdf",
        ContentFile(pdf_buffer.getvalue()),
    )

    extractor = PDFExtractor()

    text = extractor.extract(document)

    assert "Machine learning is useful." in text

@pytest.mark.django_db
def test_docx_extractor_extracts_text(
    django_user_model,
):
    user = django_user_model.objects.create_user(
        username="docxextractuser",
        password="test-password-123",
    )

    docx_buffer = BytesIO()

    docx = DocxDocument()
    docx.add_paragraph(
        "Machine learning is useful."
    )
    docx.add_paragraph(
        "RAG improves knowledge retrieval."
    )
    docx.save(docx_buffer)

    document = Document.objects.create(
        user=user,
        title="DOCX Document",
        original_filename="document.docx",
        content_type=(
            "application/"
            "vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    )

    document.file.save(
        "document.docx",
        ContentFile(docx_buffer.getvalue()),
    )

    extractor = DOCXExtractor()

    text = extractor.extract(document)

    assert "Machine learning is useful." in text
    assert "RAG improves knowledge retrieval." in text